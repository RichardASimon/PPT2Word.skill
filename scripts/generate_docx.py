#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate a professional Word document from extracted PPTX content.

Usage:
    python generate_docx.py <content.md> <slide_images.json> <media_dir> <output.docx>
    python generate_docx.py <content.md> <output.docx>  # no images

Inputs:
    content.md           — UTF-8 markdown from markitdown
    slide_images.json    — {slide_num: [filenames]} mapping (optional)
    media_dir            — directory containing extracted images
    output.docx          — output Word file path
"""

import json
import os
import re
import sys
from collections import Counter

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------- Helpers ----------

def sanitize_text(text):
    """Remove XML-incompatible control characters."""
    if not text:
        return ""
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    text = text.replace("\x0b", " ").replace("\x0c", " ")
    return text


def set_cell_shading(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, size=10, font_name="Arial"):
    text = sanitize_text(text)
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "SimSun")


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Arial"
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), "SimHei")
    if level == 0:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in h.runs:
            run.font.size = Pt(28)
            run.font.color.rgb = RGBColor(0x1E, 0x27, 0x61)
    elif level == 1:
        for run in h.runs:
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0x1E, 0x27, 0x61)
    elif level == 2:
        for run in h.runs:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x02, 0x80, 0x90)
    elif level == 3:
        for run in h.runs:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x36, 0x45, 0x4F)
    return h


def add_para(doc, text, size=10, bold=False, italic=False, align=None,
             color=None, font_name="Arial", east_asia="SimSun"):
    text = sanitize_text(text)
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east_asia)
    return p


def parse_markdown_table(table_lines):
    rows = []
    for line in table_lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if all(re.match(r"^-+\s*$", c) for c in cells):
            continue
        if cells:
            rows.append(cells)
    return rows


def add_table_to_doc(doc, rows, header_color="1E2761", alt_color="F0F4F8"):
    if not rows:
        return None
    ncols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < ncols:
            r.append("")
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    for j, cell_text in enumerate(rows[0]):
        set_cell_text(table.rows[0].cells[j], cell_text, bold=True, size=9)
        set_cell_shading(table.rows[0].cells[j], header_color)
        for run in table.rows[0].cells[j].paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for i, row in enumerate(rows[1:], 1):
        for j, cell_text in enumerate(row):
            set_cell_text(table.rows[i].cells[j], cell_text, size=9)
            if i % 2 == 0:
                set_cell_shading(table.rows[i].cells[j], alt_color)
    return table


def add_image_to_doc(doc, image_path, width_cm=14):
    if not os.path.exists(image_path):
        return False
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(image_path, width=Cm(width_cm))
        return True
    except Exception as e:
        print(f"  Warning: could not insert {image_path}: {e}", file=sys.stderr)
        return False


def add_page_number(footer):
    """Add page number field to footer."""
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text, size in [("— ", 8), ("Page ", 8), ("", 8), (" —", 8)]:
        run = fp.add_run(text)
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    # Page number field
    run3 = fp.add_run()
    run3.font.size = Pt(8)
    run3.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE   \\* MERGEFORMAT '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._r.append(fldChar1)
    run3._r.append(instrText)
    run3._r.append(fldChar2)


# ---------- Main ----------

def main():
    if len(sys.argv) < 3:
        print("Usage:")
        print("  python generate_docx.py <content.md> <slide_images.json> <media_dir> <output.docx>")
        print("  python generate_docx.py <content.md> <output.docx>")
        sys.exit(1)

    if len(sys.argv) >= 5:
        md_path = sys.argv[1]
        images_json = sys.argv[2]
        media_dir = sys.argv[3]
        output_path = sys.argv[4]
    else:
        md_path = sys.argv[1]
        images_json = None
        media_dir = None
        output_path = sys.argv[2]

    # Load inputs
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    slide_images = {}
    if images_json and os.path.exists(images_json):
        with open(images_json) as f:
            slide_images = {int(k): v for k, v in json.load(f).items()}

    # Detect part structure from markdown
    slides_raw = re.split(r"<!-- Slide number: \d+ -->", content)
    slides = [s.strip() for s in slides_raw if s.strip()]

    # Auto-detect part names from slide titles
    part_names = {}
    for slide in slides:
        lines = [l.strip() for l in slide.split("\n") if l.strip()]
        for l in lines:
            if l.startswith("!"):
                continue
            # Part dividers: "01", "02", etc.
            if re.match(r"^0[1-7]$", l):
                part_num = l
                # Next non-empty line is the part title
                idx = lines.index(l)
                for next_line in lines[idx+1:]:
                    if next_line and not next_line.startswith("!"):
                        if not re.match(r"^Part\s+\w+$", next_line):
                            part_names[part_num] = next_line
                        break
            break

    print(f"Parsed {len(slides)} slides")
    print(f"Parts detected: {part_names}")

    # Build document
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "SimSun")

    # Header/Footer
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.text = ""
    run = hp.add_run("PPT to Word Document")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    footer = section.footer
    footer.is_linked_to_previous = False
    add_page_number(footer)

    # Collect notes
    all_notes = []

    # --- Cover page ---
    for _ in range(5):
        add_para(doc, "", size=14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1E2761")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Try to extract title from first slide
    cover_title = "PPT Document"
    cover_subtitle = ""
    for slide in slides:
        lines = [l.strip() for l in slide.split("\n") if l.strip() and not l.strip().startswith("!")]
        if lines:
            cover_title = lines[0]
            if len(lines) > 1:
                cover_subtitle = lines[1]
            break

    add_para(doc, cover_title, size=32, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER,
             color=RGBColor(0x1E, 0x27, 0x61), east_asia="SimHei")
    if cover_subtitle:
        add_para(doc, "", size=10)
        add_para(doc, cover_subtitle, size=16,
                 align=WD_ALIGN_PARAGRAPH.CENTER,
                 color=RGBColor(0x02, 0x80, 0x90), east_asia="SimHei")

    for _ in range(4):
        add_para(doc, "", size=14)

    # Cover images
    if slide_images.get(1):
        for img_name in slide_images[1]:
            img_path = os.path.join(media_dir, img_name) if media_dir else img_name
            if os.path.exists(img_path):
                add_image_to_doc(doc, img_path, width_cm=4)

    doc.add_page_break()

    # --- Table of Contents ---
    add_heading(doc, "目录 / CONTENTS", level=1)
    add_para(doc, "", size=6)

    for num, title in sorted(part_names.items()):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(f"Part {num}  {title}")
        run.font.name = "Arial"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x1E, 0x27, 0x61)
        run.bold = True
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), "SimHei")

    doc.add_page_break()

    # --- Process slides ---
    for slide_idx, slide_text in enumerate(slides):
        slide_num = slide_idx + 1
        lines = slide_text.split("\n")

        # Extract notes
        notes_lines = []
        in_notes = False
        content_lines = []
        for line in lines:
            if line.strip().startswith("### Notes:"):
                in_notes = True
                continue
            if in_notes:
                notes_lines.append(line)
            else:
                content_lines.append(line)

        note_texts = []
        for nl in notes_lines:
            ns = nl.strip()
            if not ns or ns.startswith("!"):
                continue
            note_texts.append(ns)
        if note_texts:
            all_notes.append((slide_num, note_texts))

        text_lines = [l.strip() for l in content_lines
                      if l.strip() and not l.strip().startswith("!")]
        if not text_lines:
            continue

        title = text_lines[0]

        # Part divider
        if re.match(r"^0[1-7]$", title.strip()):
            part_num = title.strip()
            part_name = part_names.get(part_num, "")
            doc.add_page_break()
            add_para(doc, "", size=10)
            add_para(doc, f"Part {part_num}", size=14, italic=True,
                     align=WD_ALIGN_PARAGRAPH.CENTER,
                     color=RGBColor(0x02, 0x80, 0x90), east_asia="SimHei")
            add_heading(doc, part_name, level=1)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "12")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "028090")
            pBdr.append(bottom)
            pPr.append(pBdr)
            continue

        # Skip cover/contents
        if slide_num in (1, 2):
            continue

        # Skip English "Part One" etc
        if re.match(r"^Part\s+\w+$", title):
            continue

        # Insert images
        if slide_images.get(slide_num):
            for img_name in slide_images[slide_num]:
                img_path = os.path.join(media_dir, img_name) if media_dir else img_name
                if os.path.exists(img_path):
                    is_small = any(k in img_name for k in ['image34', 'image49', 'image46'])
                    add_image_to_doc(doc, img_path, width_cm=8 if is_small else 14)

        # Heading
        is_h2 = bool(re.search(r"[–\-—]", title)) and not title.startswith("|")
        if len(text_lines) <= 2 and not any(l.startswith("|") for l in text_lines):
            ht = re.sub(r"^.*[–\-—]\s*", "", title) if is_h2 else title
            add_heading(doc, ht, level=2)
        else:
            if is_h2:
                add_heading(doc, re.sub(r"^.*[–\-—]\s*", "", title), level=2)
            elif not title.startswith("|"):
                add_heading(doc, title, level=3)

        # Content
        table_lines = []
        para_lines = []

        def flush_table():
            nonlocal table_lines
            if table_lines:
                rows = parse_markdown_table(table_lines)
                if rows:
                    add_table_to_doc(doc, rows)
                table_lines = []

        def flush_para():
            nonlocal para_lines
            if para_lines:
                txt = "\n".join(para_lines).strip()
                for pt in txt.split("\n"):
                    pt = pt.strip()
                    if not pt:
                        continue
                    if re.search(r"Shenzhen Kaifa.*Co\., Ltd\.\s*\d*\s*$", pt):
                        continue
                    if re.match(r"^\d+$", pt):
                        continue
                    add_para(doc, pt, size=10)
                para_lines = []

        for line in content_lines:
            stripped = line.strip()
            if stripped.startswith("!"):
                continue
            if stripped.startswith("|"):
                flush_para()
                table_lines.append(stripped)
            elif stripped == "":
                if not table_lines:
                    flush_para()
            else:
                if table_lines:
                    flush_table()
                para_lines.append(stripped)

        if table_lines:
            flush_table()
        else:
            flush_para()

    # --- Notes at end ---
    if all_notes:
        doc.add_page_break()
        add_heading(doc, "资料来源 / References", level=1)
        add_para(doc, "", size=6)
        for slide_num, notes in all_notes:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(f"【Slide {slide_num}】")
            run.font.name = "Arial"
            run.font.size = Pt(9)
            run.bold = True
            run.font.color.rgb = RGBColor(0x02, 0x80, 0x90)
            for note in notes:
                color = RGBColor(0x00, 0x00, 0xCC) if note.startswith("http") else RGBColor(0x66, 0x66, 0x66)
                add_para(doc, note, size=8, color=color)

    doc.save(output_path)
    print(f"\nDocument saved: {output_path}")
    print(f"Tables: {len(doc.tables)}")
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Notes sections: {len(all_notes)}")


if __name__ == "__main__":
    main()
